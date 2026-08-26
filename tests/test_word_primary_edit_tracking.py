from __future__ import annotations

from datetime import date

from docx import Document
from pypdf import PdfReader
from PySide6.QtCore import Qt
from PySide6.QtWidgets import QInputDialog

from orthodox_calendar.calendar_engine.orthodox_calendar import OrthodoxCalendarEngine
from orthodox_calendar.models import CalendarDay, FastLevel, Fasting, Saint, ServiceRank, ServiceRankInfo, Source
from orthodox_calendar.projects import CalendarProject, ProjectSettings, ProjectStore
from orthodox_calendar.projects.model import PROJECT_SCHEMA_VERSION, saint_key
from orthodox_calendar.rendering.docx_renderer import DocxRenderer
from orthodox_calendar.rendering.pdf_renderer import PdfOptions, PdfRenderer
from orthodox_calendar.ui.calendar_editor import CalendarEditor
from orthodox_calendar.ui.main_window import DayCell


def project_with_saints() -> CalendarProject:
    days = OrthodoxCalendarEngine().generate_year(2027, "Queensland")
    target = days[6]
    target.saints = [
        Saint(1, "First", "First source saint", target.civil_date, source=Source("Authority"), display_order=0, source_order=0),
        Saint(2, "Featured", "Source featured saint", target.civil_date, source=Source("Authority"), display_order=1, source_order=1, source_primary=True),
        Saint(3, "Third", "Third source saint", target.civil_date, source=Source("Authority"), display_order=2, source_order=2),
    ]
    return CalendarProject.create("Primary Test", ProjectSettings(2027, "Queensland"), days, "2027.test")


def test_default_primary_prefers_explicit_source_designation_and_is_deterministic():
    project = project_with_saints(); first = project.resolve_days()[6]; second = project.resolve_days()[6]
    assert first.default_primary_saint_id == saint_key(first.saints[1])
    assert first.primary_saint_id == second.primary_saint_id == first.default_primary_saint_id
    assert len(first.saints) == 3 and all(item.selected for item in first.saints)


def test_default_primary_falls_back_to_authoritative_source_order():
    project = project_with_saints()
    for saint in project.source_snapshot[6]["saints"]: saint["source_primary"] = False
    project.source_snapshot[6]["default_primary_saint_id"] = ""; project.source_snapshot[6]["primary_saint_id"] = ""
    day = project.resolve_days()[6]
    assert day.primary_saint_id == saint_key(day.saints[0])


def test_noop_edit_stays_default_but_primary_change_persists_and_reset_restores(tmp_path):
    project = project_with_saints(); project.modified = False; day = project.resolve_days()[6]
    project.update_day(day, day.primary_saint_id)
    assert project.overrides == {} and not project.modified and not project.resolve_days()[6].is_edited

    new_primary = saint_key(day.saints[2]); project.update_day(day, new_primary)
    assert project.resolve_days()[6].is_edited and project.resolve_days()[6].primary_saint_id == new_primary
    path = ProjectStore(tmp_path).save(project, tmp_path / "primary.rocproject")
    reopened = ProjectStore().load(path); restored = reopened.resolve_days()[6]
    assert restored.is_edited and restored.primary_saint_id == new_primary
    reopened.reset_day(restored.civil_date); default = reopened.resolve_days()[6]
    assert not default.is_edited and default.primary_saint_id == default.default_primary_saint_id


def test_month_and_year_reset_recalculate_derived_edited_dates():
    project = project_with_saints()
    for index in (6, 14, 33):
        day = project.resolve_days()[index]; day.notes = [f"Edit {index}"]; project.update_day(day)
    assert project.edited_dates(2027, 1) == [date(2027, 1, 7), date(2027, 1, 15)]
    project.reset_month(2027, 1)
    assert project.edited_dates() == [date(2027, 2, 3)]
    project.reset_year(2027)
    assert project.edited_dates() == [] and not any(day.is_edited for day in project.resolve_days())


def test_v1_project_fixture_migrates_primary_saint_without_becoming_unreadable():
    project = ProjectStore().load(__import__("pathlib").Path(__file__).parent / "fixtures" / "sample_calendar.rocproject")
    day = project.resolve_days()[6]
    assert project.project_schema_version == PROJECT_SCHEMA_VERSION == 2
    assert day.default_primary_saint_id and day.primary_saint_id


def test_editor_explicit_primary_control_marks_only_actual_change(qtbot):
    project = project_with_saints(); day = project.resolve_days()[6]; edits = []
    editor = CalendarEditor(project.resolve_days(), on_project_edit=lambda value, primary: (project.update_day(value, primary), edits.append(primary)), initial_date=day.civil_date, source_day_provider=project.source_day)
    qtbot.addWidget(editor)
    assert editor.primary_saint.currentData() == day.default_primary_saint_id
    editor.save_overrides(); assert project.overrides == {}

    editor = CalendarEditor(project.resolve_days(), on_project_edit=lambda value, primary: project.update_day(value, primary), initial_date=day.civil_date, source_day_provider=project.source_day)
    qtbot.addWidget(editor); editor.primary_saint.setCurrentIndex(2); selected = editor.primary_saint.currentData(); editor.save_overrides()
    assert project.resolve_days()[6].primary_saint_id == selected and project.resolve_days()[6].is_edited


def test_editor_adds_secondary_saint_without_replacing_default_primary(qtbot, monkeypatch):
    project = project_with_saints(); day = project.resolve_days()[6]
    editor = CalendarEditor(project.resolve_days(), on_project_edit=lambda value, primary: project.update_day(value, primary), initial_date=day.civil_date, source_day_provider=project.source_day)
    qtbot.addWidget(editor); original_primary = editor.primary_saint.currentData()
    monkeypatch.setattr(QInputDialog, "getText", lambda *args, **kwargs: ("New Parish Saint", True))
    editor.add_saint(); assert editor.primary_saint.currentData() == original_primary
    editor.save_overrides(); restored = project.resolve_days()[6]
    assert restored.primary_saint_id == original_primary
    assert any(item.display_name == "New Parish Saint" and item.selected for item in restored.saints)


def test_pencil_indicator_tooltip_and_interaction_only_for_edited_days(qtbot):
    day = CalendarDay(date(2027, 1, 7), date(2026, 12, 25), is_edited=True)
    cell = DayCell(day); qtbot.addWidget(cell); opened = []
    cell.editRequested.connect(opened.append)
    cell.show(); assert not cell.edit_indicator.isHidden() and "Edited from default" in cell.edit_indicator.toolTip()
    qtbot.mouseDClick(cell.edit_indicator, Qt.LeftButton)
    assert opened == [day.civil_date]
    normal = DayCell(CalendarDay(date(2027, 1, 8), date(2026, 12, 26))); qtbot.addWidget(normal)
    assert normal.edit_indicator.isHidden() and "Edited from default" not in normal.toolTip()


def test_docx_is_editable_landscape_and_matches_pdf_resolved_state(tmp_path):
    project = project_with_saints(); day = project.resolve_days()[6]
    day.saints[0].selected = False; day.saints[2].display_order = 0; day.saints[1].display_order = 1
    day.primary_saint_id = saint_key(day.saints[2]); day.notes = ["PROJECT WORD NOTE"]
    day.fasting = Fasting(FastLevel.FISH, "Fish permitted", "")
    day.service_rank = ServiceRankInfo(ServiceRank.POLYELEOS, "Polyeleos", "Полиелейная служба")
    project.update_day(day, day.primary_saint_id); resolved = project.resolve_days()
    options = PdfOptions(2027, "Queensland", language="English", months=[1])
    docx_path, pdf_path = tmp_path / "calendar.docx", tmp_path / "calendar.pdf"
    DocxRenderer().render(docx_path, resolved, options); PdfRenderer().render(pdf_path, resolved, options)
    document = Document(docx_path); assert len(document.tables) == 1 and len(document.tables[0].columns) == 7
    section = document.sections[0]; assert section.page_width > section.page_height
    word_text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    pdf_text = PdfReader(pdf_path).pages[0].extract_text() or ""
    for expected in ("Third source saint", "Source featured saint", "Polyeleos", "Fish permitted", "PROJECT WORD NOTE"):
        assert expected in word_text and expected in pdf_text
    assert "First source saint" not in word_text and "First source saint" not in pdf_text


def test_russian_docx_uses_source_cyrillic_without_translation(tmp_path):
    project = project_with_saints(); day = project.resolve_days()[6]
    day.saints[0].display_name = "Святитель Николай"; day.saints[1].selected = False; day.saints[2].selected = False
    day.primary_saint_id = saint_key(day.saints[0]); project.update_day(day, day.primary_saint_id)
    output = tmp_path / "russian.docx"
    DocxRenderer().render(output, project.resolve_days(), PdfOptions(2027, "Queensland", language="Russian", months=[1]))
    text = "\n".join(cell.text for row in Document(output).tables[0].rows for cell in row.cells)
    assert "ЯНВАРЬ" not in text  # title is an editable paragraph, not duplicated into the grid
    assert "Святитель Николай" in text and "ВОСК" in text


def test_dense_docx_keeps_compact_grid_and_full_editable_details(tmp_path):
    project = project_with_saints(); day = project.resolve_days()[6]
    for index in range(4, 12):
        day.saints.append(Saint(index, f"Long {index}", f"Long additional commemoration number {index} with complete source wording", day.civil_date, display_order=index))
    project.update_day(day); output = tmp_path / "dense.docx"
    DocxRenderer().render(output, project.resolve_days(), PdfOptions(2027, "Queensland", months=[1]))
    document = Document(output); grid_text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "+9 more - see daily details" in grid_text
    assert "Long additional commemoration number 11 with complete source wording" in all_text
    assert len(document.sections) == 2
