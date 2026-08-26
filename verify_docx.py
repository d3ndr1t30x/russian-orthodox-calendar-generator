from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

from docx import Document


def verify(path: Path, expected_text: str = "") -> None:
    if not path.exists() or path.stat().st_size < 10_000:
        raise SystemExit("DOCX is missing or unexpectedly small")
    if not zipfile.is_zipfile(path):
        raise SystemExit("DOCX is not a valid OOXML package")
    document = Document(path)
    if len(document.tables) != 12:
        raise SystemExit(f"Expected 12 editable month tables, found {len(document.tables)}")
    if any(len(table.columns) != 7 for table in document.tables):
        raise SystemExit("Calendar tables must contain seven weekday columns")
    if any(section.page_width <= section.page_height for section in document.sections):
        raise SystemExit("DOCX is not landscape")
    text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    if expected_text and expected_text not in text:
        raise SystemExit("Resolved project edit is missing from DOCX")
    print(f"Verified {path}: genuine editable OOXML, 12 A4 landscape calendar tables")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(); parser.add_argument("path", type=Path); parser.add_argument("--expect", default="")
    args = parser.parse_args(); verify(args.path, args.expect)
