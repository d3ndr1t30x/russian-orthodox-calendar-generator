from __future__ import annotations

import calendar
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from orthodox_calendar import __version__
from orthodox_calendar.models import CalendarDay, FastLevel, ServiceRank
from orthodox_calendar.paths import asset_path
from orthodox_calendar.service_ranks import icon_path_for, localized_rank_name
from .pdf_renderer import IconRenderer, PdfOptions, PdfRenderer
from .publication import is_primary_saint, ordered_selected_saints


MONTHS_RU = ("", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь")
WEEKDAYS_EN = ("SUN", "MON", "TUE", "WED", "THU", "FRI", "SAT")
WEEKDAYS_RU = ("ВОСК", "ПОН", "ВТОР", "СРЕД", "ЧЕТ", "ПЯТ", "СУБ")


def _set_cell_shading(cell, colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr(); shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd"); tc_pr.append(shading)
    shading.set(qn("w:fill"), colour)


def _set_cell_margins(cell, value_dxa: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr(); margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar"); tc_pr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}"); margins.append(node)
        node.set(qn("w:w"), str(value_dxa)); node.set(qn("w:type"), "dxa")


def _prevent_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_geometry(table, total_width_dxa: int) -> None:
    table.autofit = False; column_width = total_width_dxa // 7
    tbl_pr = table._tbl.tblPr; layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW"); tbl_pr.append(width)
    width.set(qn("w:w"), str(total_width_dxa)); width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for _ in range(7):
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(column_width)); grid.append(col)
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(column_width)); tc_w.set(qn("w:type"), "dxa")


def _format_run(run, size: float, bold: bool = False, colour: str = "111111", font: str = "Arial Narrow") -> None:
    run.font.name = font; run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(colour)
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font); r_fonts.set(qn("w:hAnsi"), font); r_fonts.set(qn("w:eastAsia"), font)


def _paragraph(cell, before: float = 0, after: float = 0):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(before); paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 0.88
    return paragraph


def _compact(text: str, limit: int = 44) -> str:
    value = " ".join(text.split())
    return value if len(value) <= limit else value[:limit - 1].rstrip() + "…"


class DocxRenderer:
    """Generate an editable Word calendar directly from resolved CalendarDay data."""

    def render(self, output: Path, days: list[CalendarDay], options: PdfOptions) -> Path:
        output = Path(output); output.parent.mkdir(parents=True, exist_ok=True)
        document = Document(); section = document.sections[0]
        if options.orientation == "Landscape":
            section.orientation = WD_ORIENT.LANDSCAPE; section.page_width = Mm(297); section.page_height = Mm(210)
        else:
            section.orientation = WD_ORIENT.PORTRAIT; section.page_width = Mm(210); section.page_height = Mm(297)
        section.left_margin = section.right_margin = Mm(7)
        section.top_margin = section.bottom_margin = Mm(6)
        section.header_distance = section.footer_distance = Mm(3)
        normal = document.styles["Normal"]
        normal.font.name = "Arial Narrow"; normal.font.size = Pt(5.2)
        normal.paragraph_format.space_before = normal.paragraph_format.space_after = Pt(0)
        document.core_properties.title = f"Russian Orthodox Calendar {options.year} - {options.jurisdiction}"
        document.core_properties.author = "Russian Orthodox Calendar Generator"
        document.core_properties.comments = f"Editable calendar generated directly from resolved project data by version {__version__}."
        footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _format_run(footer.add_run(options.custom_footer or "Verify calendar data against the current official church calendar."), 4.5, colour="555555")

        usable_mm = (297 if options.orientation == "Landscape" else 210) - 14
        total_width_dxa = round(usable_mm / 25.4 * 1440)
        self._detail_days: list[CalendarDay] = []
        for page_number, month in enumerate(options.months, 1):
            self._add_month(document, [day for day in days if day.civil_date.month == month], month, options, total_width_dxa)
            if page_number < len(options.months): document.add_page_break()
        if self._detail_days:
            detail_section = document.add_section(WD_SECTION.NEW_PAGE)
            detail_section.orientation = section.orientation; detail_section.page_width = section.page_width; detail_section.page_height = section.page_height
            detail_section.left_margin = detail_section.right_margin = Mm(7); detail_section.top_margin = detail_section.bottom_margin = Mm(6)
            columns = detail_section._sectPr.xpath("./w:cols")[0]; columns.set(qn("w:num"), "3"); columns.set(qn("w:space"), "240")
            self._add_daily_details(document, options)
        document.save(output)
        return output

    def _add_month(self, document: Document, days: list[CalendarDay], month: int, options: PdfOptions, total_width_dxa: int) -> None:
        title = document.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2); title.paragraph_format.keep_with_next = True
        month_name = MONTHS_RU[month] if options.language == "Russian" else calendar.month_name[month]
        _format_run(title.add_run(month_name.upper()), 16, True, "8B1E2D", "Arial")
        _format_run(title.add_run(f"    {options.year}  |  {options.jurisdiction}"), 7.5, False, "333333", "Arial")

        weeks = calendar.Calendar(firstweekday=6).monthdayscalendar(options.year, month)
        table = document.add_table(rows=1, cols=7); table.style = "Table Grid"
        labels = WEEKDAYS_RU if options.language == "Russian" else WEEKDAYS_EN
        for index, (cell, label) in enumerate(zip(table.rows[0].cells, labels)):
            _set_cell_shading(cell, "A61E2D" if index == 0 else "28618B"); _set_cell_margins(cell, 35)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _format_run(paragraph.add_run(label), 6.5, True, "FFFFFF", "Arial")
        table.rows[0].height = Mm(6); table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST; _prevent_split(table.rows[0])

        by_number = {day.civil_date.day: day for day in days}
        row_height = max(25, int(165 / len(weeks)))
        for week in weeks:
            row = table.add_row(); row.height = Mm(row_height); row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST; _prevent_split(row)
            for number, cell in zip(week, row.cells):
                _set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if number and number in by_number: self._fill_day(cell, by_number[number], options)
        _set_table_geometry(table, total_width_dxa)

        legend = document.add_paragraph(); legend.paragraph_format.space_before = Pt(1); legend.paragraph_format.keep_with_next = False
        legend_text = "СТРОГИЙ ПОСТ | РЫБА / ВИНО / МАСЛО РАЗРЕШАЮТСЯ | ЛИТУРГИЧЕСКИЙ РАНГ" if options.language == "Russian" else "STRICT FAST | FISH / WINE / OIL PERMITTED | LITURGICAL SERVICE RANK"
        _format_run(legend.add_run(legend_text), 4.5, False, "555555", "Arial")

    def _fill_day(self, cell, day: CalendarDay, options: PdfOptions) -> None:
        state = PdfRenderer.visual_state(day)
        if state in {"great_feast", "vigil"}: _set_cell_shading(cell, "F8CACA")
        elif state == "strict_fast": _set_cell_shading(cell, "D3D3D3")
        date_line = cell.paragraphs[0]; date_line.paragraph_format.space_after = Pt(0)
        date_colour = "B00000" if day.civil_date.weekday() == 6 or state in {"great_feast", "vigil"} else "111111"
        _format_run(date_line.add_run(str(day.civil_date.day)), 10.5, True, date_colour, "Arial")
        if options.include_julian: _format_run(date_line.add_run(f"   O.S. {day.julian_date.day}"), 4.6, False, "555555")

        rank = day.service_rank.normalized_rank
        if rank not in {ServiceRank.NONE, ServiceRank.NO_DATA, ServiceRank.UNKNOWN}:
            paragraph = _paragraph(cell); path = icon_path_for(day.service_rank)
            if options.include_service_rank_icons and path and path.exists(): paragraph.add_run().add_picture(str(path), width=Mm(3.2))
            label = _compact(localized_rank_name(day.service_rank, options.language, options.rank_labels_en, options.rank_labels_ru), 32)
            _format_run(paragraph.add_run(" " + label), 4.7, True, "8B1E2D")

        saints = ordered_selected_saints(day)
        shown_feasts = day.feasts[:1]
        shown_saints = saints[:2]
        for feast in shown_feasts:
            paragraph = _paragraph(cell); major = feast.rank.value == "Great Feast" or state in {"great_feast", "vigil"}
            _format_run(paragraph.add_run(_compact(feast.name)), 5.0, major, "B00000" if major else "222222")
        for saint in shown_saints:
            paragraph = _paragraph(cell)
            _format_run(paragraph.add_run(_compact(saint.display_name)), 5.0 if is_primary_saint(day, saint) else 4.7, is_primary_saint(day, saint), "111111")
        omitted = len(day.feasts) - len(shown_feasts) + len(saints) - len(shown_saints)
        displayed_texts = [item.name for item in shown_feasts] + [item.display_name for item in shown_saints]
        needs_detail = bool(omitted or any(len(" ".join(value.split())) > 44 for value in displayed_texts)
                            or len(day.notes) > 1 or any(len(value) > 44 for value in day.notes)
                            or len(day.public_holidays) > 1 or any(len(value.name) > 44 for value in day.public_holidays))
        if needs_detail and all(item.civil_date != day.civil_date for item in self._detail_days): self._detail_days.append(day)
        if omitted:
            label = f"+{omitted} ещё - см. подробности" if options.language == "Russian" else f"+{omitted} more - see daily details"
            _format_run(_paragraph(cell).add_run(label), 4.4, False, "555555")
        if day.fasting and day.fasting.level != FastLevel.FREE:
            paragraph = _paragraph(cell); icons = [name for name in IconRenderer.permissions(day) if name != "strict_fast"]
            if options.include_fasting_icons:
                for name in icons:
                    path = asset_path("icons", f"{name}.png")
                    if path.exists(): paragraph.add_run().add_picture(str(path), width=Mm(2.8))
            fasting_text = _compact(day.fasting.period or day.fasting.detail or day.fasting.level.value)
            _format_run(paragraph.add_run((" " if icons else "") + fasting_text), 4.6, True, "333333")
        if options.include_holidays:
            for holiday in day.public_holidays[:1]:
                _format_run(_paragraph(cell).add_run(_compact(holiday.name)), 4.7, True, "243CFF")
        for note in day.notes[:1]:
            _format_run(_paragraph(cell).add_run(_compact(note)), 4.7, True, "008A18")

    def _add_daily_details(self, document: Document, options: PdfOptions) -> None:
        title = document.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _format_run(title.add_run("ПОДРОБНОСТИ ПО ДНЯМ" if options.language == "Russian" else "DAILY CALENDAR DETAILS"), 15, True, "8B1E2D", "Arial")
        intro = document.add_paragraph()
        _format_run(intro.add_run("Полные редактируемые списки памятей, не поместившиеся в компактную месячную сетку." if options.language == "Russian" else "Complete editable commemoration lists that did not fit in the compact monthly grid."), 8, False, "555555", "Arial")
        for day in self._detail_days:
            paragraph = document.add_paragraph(); paragraph.paragraph_format.space_before = Pt(1.5); paragraph.paragraph_format.space_after = Pt(0); paragraph.paragraph_format.line_spacing = 0.82
            date_label = (f"{day.civil_date.day} {MONTHS_RU[day.civil_date.month].lower()} {day.civil_date.year}" if options.language == "Russian" else day.civil_date.strftime("%d %B %Y"))
            _format_run(paragraph.add_run(date_label), 6.2, True, "8B1E2D", "Arial Narrow")
            rank = localized_rank_name(day.service_rank, options.language, options.rank_labels_en, options.rank_labels_ru)
            if rank: _format_run(paragraph.add_run(" | " + rank), 5.2, True, "555555", "Arial Narrow")
            separator = " — "
            for feast in day.feasts:
                _format_run(paragraph.add_run(separator + feast.name), 5.3, True, "B00000", "Arial Narrow"); separator = "; "
            for saint in ordered_selected_saints(day):
                _format_run(paragraph.add_run(separator + saint.display_name), 5.2, is_primary_saint(day, saint), "111111", "Arial Narrow"); separator = "; "
            if day.fasting and day.fasting.level != FastLevel.FREE:
                _format_run(paragraph.add_run(" | " + (day.fasting.period or day.fasting.detail or day.fasting.level.value)), 5.1, True, "444444", "Arial Narrow")
            for holiday in day.public_holidays:
                _format_run(paragraph.add_run(" | " + holiday.name), 5.1, True, "243CFF", "Arial Narrow")
            for note in day.notes:
                _format_run(paragraph.add_run(" | " + note), 5.1, True, "008A18", "Arial Narrow")
